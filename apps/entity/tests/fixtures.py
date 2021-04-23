import os
from django.conf import settings
from django.core.files.uploadedfile import InMemoryUploadedFile
from docx import Document as Docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from fpdf import FPDF


class ActFileFixture:
    """
    ActFileFixture
    Abstract class to create files fixtures. Generic functions should be
    implemented in this file.
    :param file_name: A string representing the name of a file
    """

    def __init__(self, file_name):
        self.file_name = file_name

    def build_in_memory_uploaded_file(self, opened_file, file_name="test_file", file_field="file"):
        """
        build_in_memory_uploaded_file
        Given an opened file, returns an InMemoryUploadedFile instance of that
        file. This is generally used to test attached files for an endpoint.
        :param opened_file: An opened file using the open(*args) function
        :param file_name: An optional string representing a file name
        :param file_field: An optiona string representing a fiel field
        """
        return InMemoryUploadedFile(
            opened_file.buffer, file_field, file_name, None, None, opened_file.buffer.tell(), None
        )

    def create_dummy_sentence(self):
        header_paragraphs = [
            "2019 - Año del 25º Aniversario del reconocimiento de la autonomía de la Ciudad de Buenos Aires",
            "JUZGADO DE 1RA INSTANCIA EN LO PENAL CONTRAVENCIONAL Y DE FALTAS N°10   SECRETARÍA UNICA",
            "Ordena trabar embargo sobre cuenta bancaria",
            "CUIJ: IPP J-01-00006472-1/2021-0",
            "Actuación Nro: 18236125/2021",
        ]
        footer_paragraphs = [
            "Juzgado PCyF Nº 10 - Tacuarí 138, 7º Piso - juzcyf10@jusbaires.gob.ar - 4014-6821/20 - @jpcyf10"
        ]
        topic_section = (
            {
                "heading": "ORDENA EMBARGO PREVENTIVO\n",
                "paragraphs": ["Ordena trabar embargo sobre cuenta bancaria"],
                "date_and_place": "Buenos Aires, 18 de febrero de 2021.",
            },
        )
        body_sections = [
            {
                "heading": "ANTECEDENTES\n",
                "paragraphs": [
                    "    El pasado 25 de enero de 2021 decidí revocar la suspensión del proceso a prueba concedida al acusado el 03/12/2020 en el marco de este caso penal.",
                    "    En razón de ello regulé alimentos provisorios en favor de Juan Lopez, Ana Lopez , Rosario Lopez Y Agusto Lopez domiciliados junto a su madre Roberta Paz , por el monto de dos mil pesos ($2.000) mensuales, a partir del mes de Marzo y hasta tanto la justicia civil regule el régimen definitivo.\n",
                    "    Esta decisión fue apelada por la defensa y finalmente el pasado 05/02/2021 la Cámara de Apelaciones la confirmó en su totalidad, venciendo el plazo para presentar recursos sin que la defensa efectuara una nueva presentación.",
                    "    En función de ello, el 10/02/2021 decidí intimar al acusado a través de su defensa a fin de que acredite el cumplimiento de las obligaciones mencionadas.",
                    "    En lo que aquí interesa, tambien hay que tener en cuenta que el acusado de  nacionalidad argentina, con numero de documento 11.111.111, propietario de un automovil  de marca Ford Fiesta con patente AB123CD , radicado en la dirección Chiclana 1429 , Bernal, Provincia de Buenos Aires desde hace 3 años. Las pruebas que se remiten son imagenAuto1.jpg  y videoAuto.mp4.\n",
                    "    En virtud de ello, solicitó al juzgado que instrumentara las medidas necesarias para que los fondos existentes en la mencionada cuenta bancaria fueran transferidos a una cuenta bancaria a nombre de la denunciante, quien convive con los niños y niñas beneficiarios del subsidio. ",
                    "    Explicó que el pedido se basaba en la inexistencia de una tarjeta de débito por parte de su asistido que le permitiera realizar la transferencia por cajero automático o por ventanilla desde cualquier sucursal de la entidad bancaria señalada, así como su incapacidad manifiesta para operar a través de los canales bancarios y electrónicos habituales para el común de las personas.",
                    "    En oportunidad de expedirse, la Fiscalía refirió que estaba realizando las tareas necesarias a fin de poder retomar contacto con la víctima ya que los teléfonos aportados en la investigación se encuentran fuera del área de cobertura.",
                    "    En función de ello e invocando la Convención de los Derechos del Niño manifestó que la única forma de asegurar las condiciones de vida necesarias para el desarrollo de los niños y niñas en este caso es trabando embargo en la cuenta bancaria sobre la suma existente y sobre las futuras sumas de dinero que mensualmente se acrediten en dicha cuenta, hasta tanto pueda contactarse con la denunciante.\n",
                ],
            },
            {
                "heading": "ARGUMENTOS\n",
                "paragraphs": [
                    "   El pedido de embargo busca garantizar que las sumas de dinero entregadas por el Estado para la manutención y el cuidado del niño y las niñas presentes en este caso sean efectivamente utilizadas a dichos fines.",
                    "   En este sentido, ya sostuve en oportunidades anteriores que las hipótesis según las cuales considero que resulta procedente el embargo, son la necesidad de asegurar el cumplimiento de la pena pecuniaria que podría imponerse al acusado en caso de recaer condena, como así también la posibilidad de garantizar la reparación del daño causado por el delito (art. 188 CPP) o la indemnización civil derivada de la deuda alimentaria (art. 343 CPP).\n",
                    "   En segundo lugar, se encuentra configurado el requisito relativo al mérito sustantivo, dado que de acuerdo con la prueba señalada por la señora Fiscal en su requerimiento de juicio, se encuentran reunidos elementos suficientes como sostener la adecuada fundamentación de la acusación.",
                    "   A ello debe sumarse que ya analicé previamente los antecedentes del caso y los elementos de cargo aportados por la fiscalía y en función de ello decidí obligar al acusado a que entregase a la denunciante la totalidad del dinero que reciba de ANSES en concepto de asignación por sus hijos e hijas.",
                    "   Efectivamente, de la prueba aportada por la Fiscalía surge que es la denunciante quien se encuentra al cuidado exclusivo de los hijos que tiene en común con el acusado.",
                    "   En tercer lugar, se encuentran suficientemente comprobados los riesgos procesales alegados por la señora Fiscal, dado que se acreditó a partir de los elementos probatorios aportados en el caso, que dicha cuenta es de exclusiva titularidad del acusado, por lo que es necesario reducir el riesgo de que pudiera disponer libremente de dicho dinero para evadir el cumplimiento de la eventual pena pecuniaria y de los demás deberes resarcitorios que pudieran surgir. ",
                    "   Por lo expuesto, entiendo que la medida cautelar solicitada resulta procedente para garantizar el interés superior de los niños que resultarían víctimas del delito imputado al acusado, como así también los derechos de la denunciante.",
                    "   Concretamente, voy a disponer que se trabe embargo sobre la cuenta de caja de ahorro social nro. 234521 del Banco Columbia , abierta en la Sucursal nro. 22 de la localidad de Quilmes de la Provincia de Buenos Aires, a nombre de Carlos Lopez , DNI 62.122.522, asociada a la casilla de mail carloslopes@mail.com.ar.Ademas las sumas que a futuro perciba allí el acusado en representación de sus hijos menores de edad, haciéndole saber al Banco Columbia que deberá informar mensualmente los depósitos realizados a la Fiscalía interviniente, ya que mi jurisdicción sobre el caso cesará con la remisión del caso a juicio, en virtud del estadio procesal en que nos encontramos.",
                ],
            },
        ]
        return {
            "header_paragraphs": header_paragraphs,
            "footer_paragraphs": footer_paragraphs,
            "topic_section": topic_section,
            "body_sections": body_sections,
        }


class ActPdfFileFixture(ActFileFixture):
    """
    ActPdfFileFixture
    Given a file name and structured data, returns a pdf file object.
    :param file_name: A string representing a file name
    :param data: a dictionary representing data to create a document. See the
    create_dummy_sentence function.
    """

    def __init__(self, file_name, data=None):
        super().__init__(file_name)
        if data is None:
            data = self.create_dummy_sentence()
        document = PdfDocument(file_name, data=data)
        document = document.build()
        self.output_dir = document.create()


class ActDocxFileFixture(ActFileFixture):
    def __init__(self, file_name, data=None):
        super().__init__(file_name)
        if data is None:
            data = self.create_dummy_sentence()
        document = DocxDocument(file_name, data=data)
        document = document.build()
        self.output_dir = document.create()


class PdfDocument(FPDF):
    """
    PdfDocument is responsible for creating pdf documents in run time.
    :param file_name: a string representing a file name
    :param data: a dictionary representing data to create a document. See the
    create_dummy_sentence function.
    """

    def __init__(self, file_name, *, font_family="Times", data):
        super().__init__()
        self.file_name = file_name
        self.font_family = font_family
        self.header_paragraphs = data["header_paragraphs"]
        self.footer_paragraphs = data["footer_paragraphs"]
        self.topic_sections = data["topic_section"][0]
        self.body_sections = data["body_sections"]

    def header(self):
        """
        header
        Overrides FPDF empty header implementation. Creates a header everytime
        add_page() is called
        """
        if self.page == 1:
            self.set_font(self.font_family, "", 10)
            for index, paragraph in enumerate(self.header_paragraphs):
                if index == 0:
                    self.cell(w=0, h=10, txt=paragraph, border=0, ln=1, align="C")
                else:
                    self.set_font(self.font_family, "B", 10)
                    self.cell(w=0, h=5, txt=paragraph, border=0, ln=1, align="R")

    def footer(self):
        """
        footer
        Overrides FPDF empty footer implementation. Creates a footer everytime
        add_page() is called
        """
        self.set_font(self.font_family, "", 10)
        self.ln(h=10)
        for paragraph in self.footer_paragraphs:
            self.cell(w=0, h=10, txt=paragraph, border=0, ln=1, align="C")

    def build_topic_section(self):
        """
        build_topic_section
        Builds a heading, a subheading as pararaphs and a paragraph to the
        document
        """
        self.set_font(self.font_family, "BU", 16)
        self.cell(w=0, h=12, txt=self.topic_sections["heading"], ln=1, align="C")
        for paragraph in self.topic_sections["paragraphs"]:
            self.set_font(self.font_family, "BIU", 12)
            self.cell(w=0, h=8, txt=paragraph, ln=1, align="L")
        self.set_font(self.font_family, "", 12)
        self.write(h=8, txt=self.topic_sections["date_and_place"])

    def build_body_section(self):
        """
        build_body_section
        Builds body paragraphs to the document
        """
        self.ln(10)
        for body_section in self.body_sections:
            self.set_font(self.font_family, "BU", 16)
            self.write(h=12, txt=body_section["heading"])
            for paragraph in body_section["paragraphs"]:
                self.set_font(self.font_family, "", 12)
                self.write(h=8, txt=paragraph)

    def build(self):
        """
        build
        Builds all document sections
        """
        self.add_page()
        self.build_topic_section()
        self.build_body_section()
        return self

    def create(self):
        """
        create
        This method expects build to have been called. Creates a file with the
        given file_name in the current directory and returns the output path.
        """
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), self.file_name)
        self.output(output_dir, "F")
        return output_dir


class DocxDocument:
    """
    DocxDocument is responsible for creating docx documents in run time.
    :param file_name: a string representing a file name
    :param data: a dictionary representing data to create a document. See the
    create_dummy_sentence function.
    """

    def __init__(self, file_name, *, data):
        self.file_name = file_name
        self.document = Docx()
        section = self.document.sections[0]
        section.different_first_page_header_footer = True
        self.first_page_header = section.first_page_header
        self.first_page_footer = section.first_page_footer
        self.footer = section.footer
        self.header_paragraphs = data["header_paragraphs"]
        self.footer_paragraphs = data["footer_paragraphs"]
        self.topic_sections = data["topic_section"][0]
        self.body_sections = data["body_sections"]

    def build_header(self):
        """
        build_header
        Builds document headers
        """
        for index, paragraph in enumerate(self.header_paragraphs):
            if index == 0:
                self.first_page_header.add_paragraph(paragraph, style="Body Text 3")
            else:
                p = self.first_page_header.add_paragraph(style="Body Text 3")
                p.add_run(paragraph).bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def build_footer(self):
        """
        build_footer
        Builds document footers
        """
        for paragraph in self.footer_paragraphs:
            self.first_page_footer.add_paragraph(paragraph, style="Body Text 3")
            self.footer.add_paragraph(paragraph, style="Body Text 3")

    def build_topic_section(self):
        h = self.document.add_heading("", 1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h_run = h.add_run(self.topic_sections["heading"])
        h_run.underline = WD_UNDERLINE.SINGLE
        for paragraph in self.topic_sections["paragraphs"]:
            p = self.document.add_paragraph(style="Body Text 2")
            p_run = p.add_run(paragraph)
            p_run.bold = True
            p_run.italic = True
            p_run.underline = WD_UNDERLINE.SINGLE
        self.document.add_paragraph(self.topic_sections["date_and_place"], style="Body Text 2")

    def build_body_section(self):
        for body_section in self.body_sections:
            h = self.document.add_heading("", 1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            h_run = h.add_run(body_section["heading"])
            h_run.underline = WD_UNDERLINE.SINGLE
            [self.document.add_paragraph(paragraph, style="Body Text 2") for paragraph in body_section["paragraphs"]]

    def build(self):
        """
        build
        Builds all document sections
        """
        self.build_header()
        self.build_footer()
        self.build_topic_section()
        self.build_body_section()
        return self

    def create(self):
        """
        This method expects build to have been called. Creates a file with the
        given file_name in the current directory and returns the output path.
        """
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), self.file_name)
        self.document.save(output_dir)
        return output_dir
